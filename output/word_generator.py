# output/word_generator.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import FicheEntretien
from datetime import date
import io


def generate_word_doc(fiche: FicheEntretien, output) -> None:
    """
    Génère la fiche d'entretien Word.
    output : chemin fichier (str) ou buffer BytesIO
    """
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── En-tête ───────────────────────────────────────────────────────────────
    titre = doc.add_heading(f"Entretien bilan — {fiche.client_exercice}", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_color(titre.runs[0], 0x0F, 0x20, 0x44)

    sub = doc.add_paragraph(f"Fiche préparée le {date.today().strftime('%d/%m/%Y')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x7A, 0x99)
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Synthèse exécutive ────────────────────────────────────────────────────
    doc.add_heading("Synthèse", level=1)
    doc.add_paragraph(fiche.synthese_executive)

    # ── Points de vigilance ───────────────────────────────────────────────────
    if fiche.points_vigilance:
        doc.add_heading("⚠ Points de vigilance", level=1)
        for point in fiche.points_vigilance:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(point)
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # ── Plan d'entretien ──────────────────────────────────────────────────────
    doc.add_heading("Plan d'entretien", level=1)
    for pt in fiche.plan_entretien:
        doc.add_heading(f"{pt.ordre}. {pt.theme}", level=2)

        p1 = doc.add_paragraph()
        p1.add_run("Contexte : ").bold = True
        p1.add_run(pt.contexte_chiffre)

        p2 = doc.add_paragraph()
        p2.add_run("Question : ").bold = True
        r = p2.add_run(pt.question_ouverte)
        r.italic = True

        if pt.mission_associee:
            p3 = doc.add_paragraph()
            p3.add_run("→ Mission associée : ").bold = True
            p3.add_run(pt.mission_associee)

    # ── Missions à proposer ───────────────────────────────────────────────────
    doc.add_heading("Missions à proposer", level=1)
    for m in fiche.missions_a_proposer:
        doc.add_heading(m.get("titre", "—"), level=2)

        p_arg = doc.add_paragraph()
        p_arg.add_run("Argumentaire : ").bold = True
        p_arg.add_run(m.get("argumentaire_personnalise", ""))

        p_ben = doc.add_paragraph()
        p_ben.add_run("Bénéfice attendu : ").bold = True
        p_ben.add_run(m.get("benefice_attendu", ""))

        urgence = m.get("urgence", "")
        color   = RGBColor(0xC0, 0x00, 0x00) if urgence == "immédiate" else \
                  RGBColor(0xE6, 0x7E, 0x22) if urgence == "court terme" else \
                  RGBColor(0x22, 0x55, 0xA4)
        p_urg = doc.add_paragraph()
        p_urg.add_run("Urgence : ").bold = True
        r_urg = p_urg.add_run(urgence.upper())
        r_urg.bold = True
        r_urg.font.color.rgb = color

    # ── Éléments à recueillir ─────────────────────────────────────────────────
    if fiche.elements_a_recueillir:
        doc.add_heading("Éléments à recueillir lors du RDV", level=1)
        for elem in fiche.elements_a_recueillir:
            doc.add_paragraph(elem, style="List Bullet")

    # ── Conclusion ────────────────────────────────────────────────────────────
    doc.add_heading("Comment conclure le rendez-vous", level=1)
    p_ccl = doc.add_paragraph(fiche.conclusion_conseillee)
    p_ccl.runs[0].italic = True

    # Sauvegarde
    if isinstance(output, (str,)):
        doc.save(output)
    else:
        doc.save(output)


def _set_color(run, r, g, b):
    run.font.color.rgb = RGBColor(r, g, b)
